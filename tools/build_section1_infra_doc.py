from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\Nguend Arthur Johann\Desktop\borderwatch")
OUT = ROOT / "Section1_Infrastructure_Setup_UPDATED.docx"
DIAGRAM = ROOT / "Section1_Infrastructure_Diagram.png"


BLUE = "1F4E79"
TEAL = "0F766E"
GRAY = "E7EEF6"
LIGHT = "F7FAFC"
DARK = "263238"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = Inches(0.04)
            cell.margin_bottom = Inches(0.04)
            cell.margin_left = Inches(0.06)
            cell.margin_right = Inches(0.06)
            if header and row_idx == 0:
                shade(cell, BLUE)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.color.rgb = RGBColor.from_string(BLUE)
    else:
        r.font.color.rgb = RGBColor.from_string(TEAL)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "EAF7F4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def draw_box(draw, xy, title, subtitle, fill, outline):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=3)
    try:
        font_title = ImageFont.truetype("arial.ttf", 18)
        font_sub = ImageFont.truetype("arial.ttf", 13)
    except Exception:
        font_title = ImageFont.load_default()
        font_sub = ImageFont.load_default()
    tw = draw.textbbox((0, 0), title, font=font_title)[2]
    sw = draw.textbbox((0, 0), subtitle, font=font_sub)[2]
    draw.text((x1 + (x2 - x1 - tw) / 2, y1 + 18), title, fill="#" + DARK, font=font_title)
    draw.text((x1 + (x2 - x1 - sw) / 2, y1 + 48), subtitle, fill="#4A5568", font=font_sub)


def draw_arrow(draw, start, end, color="#334155"):
    draw.line([start, end], fill=color, width=3)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)]
    else:
        pts = [(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)]
    draw.polygon(pts, fill=color)


def make_diagram():
    img = Image.new("RGB", (1400, 780), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arial.ttf", 34)
        font_small = ImageFont.truetype("arial.ttf", 16)
    except Exception:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()
    draw.text((50, 30), "BorderWatch Infrastructure Deployment Topology", fill="#" + BLUE, font=font_title)
    draw.text((52, 78), "Docker Compose environment with Nginx reverse proxy, Spring services, RabbitMQ, MySQL, Prometheus and Grafana", fill="#475569", font=font_small)

    draw_box(draw, (70, 140, 300, 230), "GPS Beacons", "POST /api/v1/telemetry", "#E0F2FE", "#0284C7")
    draw_box(draw, (380, 140, 610, 230), "React SPA", "served by Nginx", "#DCFCE7", "#16A34A")
    draw_box(draw, (720, 140, 1030, 230), "Nginx Gateway", "port 80, rate limits, routing", "#F0F9FF", "#1D4ED8")

    draw_box(draw, (120, 330, 350, 430), "Ingestion x2", "8081, auth + telemetry", "#FFF7ED", "#EA580C")
    draw_box(draw, (430, 330, 660, 430), "Tracking x2", "8082, geofencing", "#ECFDF5", "#059669")
    draw_box(draw, (740, 330, 980, 430), "Compliance x2", "8083, violations + audit", "#F5F3FF", "#7C3AED")
    draw_box(draw, (1080, 330, 1310, 430), "RabbitMQ", "AMQP + management UI", "#FEF3C7", "#D97706")

    draw_box(draw, (170, 560, 400, 650), "MySQL 8", "schema + seed data", "#EFF6FF", "#2563EB")
    draw_box(draw, (520, 560, 750, 650), "Prometheus", "scrapes /actuator", "#FFF1F2", "#E11D48")
    draw_box(draw, (850, 560, 1080, 650), "Grafana", "datasource on :3000", "#FFF7ED", "#F97316")
    draw_box(draw, (1130, 560, 1340, 650), "Alertmanager", "alerts on :9093", "#F8FAFC", "#64748B")

    draw_arrow(draw, (300, 185), (720, 185), "#1D4ED8")
    draw_arrow(draw, (610, 185), (720, 185), "#1D4ED8")
    draw_arrow(draw, (850, 230), (240, 330), "#EA580C")
    draw_arrow(draw, (880, 230), (545, 330), "#059669")
    draw_arrow(draw, (910, 230), (860, 330), "#7C3AED")
    draw_arrow(draw, (350, 380), (1080, 380), "#D97706")
    draw_arrow(draw, (660, 380), (1080, 380), "#D97706")
    draw_arrow(draw, (860, 430), (285, 560), "#2563EB")
    draw_arrow(draw, (240, 430), (285, 560), "#2563EB")
    draw_arrow(draw, (635, 560), (635, 430), "#E11D48")
    draw_arrow(draw, (750, 605), (850, 605), "#F97316")
    draw_arrow(draw, (1080, 605), (1130, 605), "#64748B")

    draw.rectangle((50, 700, 1350, 730), fill="#F8FAFC", outline="#CBD5E1")
    draw.text((70, 706), "Deployment command: ./scripts/start.sh -> docker compose up -d --build. Public entry point: http://localhost:80 or VPS_PUBLIC_IP:80.", fill="#334155", font=font_small)
    img.save(DIAGRAM)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "111827")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(229, 231, 235)
        if idx != len(lines) - 1:
            r.add_break()
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=8.5)
        shade(table.rows[0].cells[i], BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.5)
    style_table(table, header=False)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def build_doc():
    make_diagram()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].paragraph_format.space_before = Pt(10)
        styles[name].paragraph_format.space_after = Pt(5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BORDERWATCH")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Section 1 - Infrastructure Setup")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("High-Availability Goods-in-Transit Surveillance System")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(DARK)

    add_table(
        doc,
        ["Project", "Section", "Deployment Model", "Evidence"],
        [[
            "BorderWatch distributed customs surveillance",
            "Infrastructure Setup (15 marks)",
            "VPS-ready Docker Compose stack",
            "docker-compose.yml, Dockerfiles, Nginx, scripts, monitoring configs",
        ]],
        [1.7, 1.55, 1.8, 2.7],
    )

    add_callout(
        doc,
        "Assessment focus",
        "This section documents the real infrastructure implemented in the project repository: containerized Spring Boot services, a React frontend served through Nginx, RabbitMQ messaging, MySQL persistence, Prometheus/Grafana monitoring, and Bash automation scripts. Terraform and Kubernetes are optional extensions; the current project deploys through Docker Compose.",
    )

    add_heading(doc, "1. Infrastructure Objective", 1)
    doc.add_paragraph(
        "BorderWatch requires an infrastructure capable of receiving GPS telemetry from truck beacons, routing API traffic to the correct backend service, detecting corridor deviations, recording compliance violations, and exposing monitoring data for operations. The implemented environment is designed for local demonstration and VPS deployment using the same Docker Compose topology."
    )
    add_bullets(
        doc,
        [
            "Provision a Linux host or on-premise server with Docker Engine and Docker Compose.",
            "Run three Spring Boot microservices with two replicas each for ingestion, tracking, and compliance.",
            "Expose a single public HTTP entry point through Nginx while keeping internal services on the Docker bridge network.",
            "Use RabbitMQ for asynchronous telemetry and violation event delivery.",
            "Use MySQL 8 for users, trucks, telemetry, violations, audit chain, notifications, and seed data.",
            "Provide operational visibility through Prometheus, Grafana, Alertmanager, and node-exporter.",
        ],
    )

    add_heading(doc, "2. Provisioned Resources", 1)
    add_table(
        doc,
        ["Resource", "Implemented choice", "Project evidence", "Purpose"],
        [
            ["Compute host", "Linux/VPS-ready host", "scripts/start.sh, docker-compose.yml", "Runs the complete BorderWatch stack from one command."],
            ["Container runtime", "Docker + Docker Compose", "docker-compose.yml", "Builds and orchestrates 13 service containers."],
            ["Application gateway", "Nginx 1.27 Alpine", "gateway/Dockerfile, gateway/nginx.conf", "Serves React assets and reverse-proxies API traffic."],
            ["Database", "MySQL 8.0", "scripts/schema.sql, scripts/seed.sql", "Stores operational data and seeded demo users/trucks."],
            ["Message broker", "RabbitMQ 3.13 management image", "docker-compose.yml, RabbitMqConfig.java", "Decouples telemetry ingestion from tracking and compliance."],
            ["Monitoring", "Prometheus, Grafana, Alertmanager, node-exporter", "monitoring/*", "Collects service metrics and infrastructure health signals."],
        ],
        [1.35, 1.4, 2.0, 2.4],
    )

    add_heading(doc, "3. Container Topology", 1)
    doc.add_paragraph(
        "The Docker Compose file defines a bridge network named bw-net and persistent volumes for MySQL, RabbitMQ, Prometheus, and Grafana. Application services are built from project Dockerfiles and are exposed internally to Nginx, not directly to the public network."
    )
    add_table(
        doc,
        ["Service group", "Replicas", "Internal port", "Main responsibility"],
        [
            ["ingestion-service", "2", "8081", "Authentication, user/truck APIs, telemetry ingestion."],
            ["tracking-engine", "2", "8082", "Geofence analysis and OFF_ROUTE event generation."],
            ["compliance-service", "2", "8083", "Violation persistence, penalty classification, audit chain API."],
            ["gateway", "1", "80", "React SPA hosting, reverse proxy, API rate limiting."],
            ["mysql", "1", "3306", "Relational persistence and seed data loading."],
            ["rabbitmq", "1", "5672 / 15672", "AMQP queues and management UI."],
            ["prometheus/grafana/alertmanager/node-exporter", "4", "9090 / 3000 / 9093 / 9100", "Monitoring and alerting stack."],
        ],
        [1.8, 0.8, 1.2, 3.5],
    )

    add_heading(doc, "4. Infrastructure Diagram", 1)
    doc.add_paragraph(
        "The diagram below summarizes the implemented deployment path from users and GPS beacons through the reverse proxy, backend services, broker, database, and monitoring layer."
    )
    doc.add_picture(str(DIAGRAM), width=Inches(7.2))
    cap = doc.add_paragraph("Figure 1. BorderWatch Docker Compose infrastructure topology.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8.5)
    cap.runs[0].italic = True

    add_heading(doc, "5. Networking, Firewall and Reverse Proxy", 1)
    doc.add_paragraph(
        "Networking is handled in layers. Docker Compose isolates backend containers on bw-net, Nginx provides the single application ingress point, and host/VPS firewall rules should expose only the ports required for users, administrators, and monitoring."
    )
    add_table(
        doc,
        ["Port", "Service", "Exposure", "Reason"],
        [
            ["80", "Nginx gateway", "Public", "Main entry point for React SPA and /api/v1 routes."],
            ["3000", "Grafana", "Admin only", "Dashboard access; should be restricted on a real VPS."],
            ["9090", "Prometheus", "Admin only", "Metrics query UI; should not be public in production."],
            ["9093", "Alertmanager", "Admin only", "Alert routing UI; should be restricted."],
            ["15672", "RabbitMQ management", "Admin only", "Broker management; never expose broadly."],
            ["13306", "MySQL host mapping", "Local/dev only", "External database access for development; close on production VPS."],
            ["22", "SSH", "Admin only", "VPS administration; restrict by source IP."],
        ],
        [0.75, 1.5, 1.05, 3.85],
    )
    doc.add_paragraph("Recommended VPS firewall commands:")
    add_code_block(
        doc,
        [
            "sudo ufw default deny incoming",
            "sudo ufw default allow outgoing",
            "sudo ufw allow 22/tcp",
            "sudo ufw allow 80/tcp",
            "sudo ufw allow 443/tcp   # if TLS is enabled with a certificate",
            "sudo ufw allow from <ADMIN_IP> to any port 3000,9090,9093,15672 proto tcp",
            "sudo ufw enable",
        ],
    )

    add_heading(doc, "6. Reverse Proxy Routing", 1)
    doc.add_paragraph(
        "The Nginx gateway uses least-connection upstreams for the replicated Spring services and applies separate request limits for authentication and telemetry traffic. React assets are served from /usr/share/nginx/html, while API paths are routed to the correct internal service."
    )
    add_table(
        doc,
        ["Route", "Upstream", "Configuration evidence"],
        [
            ["/api/v1/auth", "ingestion-1/ingestion-2:8081", "gateway/nginx.conf rate-limits auth at 10 r/s."],
            ["/api/v1/telemetry", "ingestion-1/ingestion-2:8081", "gateway/nginx.conf allows higher telemetry burst capacity."],
            ["/api/v1/trucks, /api/v1/users", "ingestion service", "Routes registry and user APIs to ingestion."],
            ["/api/v1/tracking", "tracking-1/tracking-2:8082", "Routes geofence analysis to tracking engine."],
            ["/api/v1/compliance", "compliance-1/compliance-2:8083", "Routes violations, stats, and chain APIs to compliance."],
            ["/", "React SPA", "Nginx try_files fallback supports client-side routing."],
        ],
        [1.7, 2.0, 3.1],
    )

    add_heading(doc, "7. Installed Technologies", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Installed/configured through"],
        [
            ["Frontend", "React 18, TypeScript, Vite, Zustand, React-Leaflet, Recharts", "frontend-react/package.json and gateway Dockerfile."],
            ["Backend", "Java 21, Spring Boot 3.3, Spring Web, Security, JPA, AMQP, Actuator", "Three Maven POM files and service Dockerfiles."],
            ["Messaging", "RabbitMQ 3.13 Management", "docker-compose.yml and RabbitMqConfig classes."],
            ["Database", "MySQL 8.0", "docker-compose.yml, scripts/schema.sql, scripts/seed.sql."],
            ["Gateway", "Nginx 1.27 Alpine", "gateway/Dockerfile and gateway/nginx.conf."],
            ["Observability", "Prometheus, Grafana, Alertmanager, node-exporter", "monitoring directory and compose services."],
            ["Automation", "Bash scripts and Jenkinsfile", "scripts/start.sh, stop.sh, run-tests.sh, chaos_monkey.sh, Jenkinsfile."],
        ],
        [1.2, 2.8, 3.0],
    )

    add_heading(doc, "8. Deployment Automation Scripts", 1)
    doc.add_paragraph(
        "The project uses Bash automation rather than Terraform. The scripts are simple, auditable, and sufficient for a VPS or lab machine where Docker is already installed."
    )
    add_table(
        doc,
        ["Script", "Function", "How it supports infrastructure marks"],
        [
            ["scripts/start.sh", "Builds and starts all containers with docker compose up -d --build.", "Demonstrates repeatable environment setup."],
            ["scripts/stop.sh", "Stops the stack with docker compose down.", "Provides controlled shutdown."],
            ["scripts/run-tests.sh", "Runs Maven tests for the three Spring services.", "Supports validation before deployment."],
            ["scripts/chaos_monkey.sh", "Kills selected containers and measures recovery time.", "Demonstrates resilience and self-healing behavior."],
            ["scripts/schema.sql / seed.sql", "Creates database schema and demo data.", "Automates database provisioning."],
        ],
        [1.5, 2.6, 3.0],
    )
    doc.add_paragraph("Core deployment sequence:")
    add_code_block(
        doc,
        [
            "cp .env.example .env",
            "# edit .env: MYSQL_ROOT_PASSWORD, MYSQL_PASSWORD, RABBITMQ_PASSWORD, JWT_SECRET, GRAFANA_PASSWORD",
            "./scripts/start.sh",
            "docker compose ps",
            "curl http://localhost/health",
        ],
    )

    add_heading(doc, "9. Monitoring and Health Checks", 1)
    doc.add_paragraph(
        "Each Spring service exposes health and Prometheus endpoints through Spring Actuator. Docker Compose uses health checks before allowing dependent services such as Nginx to start. Prometheus scrapes ingestion, tracking, compliance, and node-exporter targets."
    )
    add_table(
        doc,
        ["Monitoring item", "Implemented configuration", "Purpose"],
        [
            ["Spring Actuator", "health,prometheus,info endpoints", "Service liveness and metrics export."],
            ["Prometheus scrape jobs", "ingestion, tracking, compliance, node-exporter", "Central metrics collection."],
            ["Alert rules", "ServiceDown, HighCPU, LowMemory", "Operational alerting for outages and resource pressure."],
            ["Grafana datasource", "Prometheus provisioned as default", "Dashboard-ready monitoring datasource."],
            ["Docker health checks", "wget /actuator/health for each service", "Container readiness and dependency ordering."],
        ],
        [1.55, 2.65, 2.9],
    )

    add_heading(doc, "10. Deployment Evidence Links and Screenshots", 1)
    doc.add_paragraph(
        "The following URLs are the expected evidence points after running ./scripts/start.sh on a lab machine or VPS. These are the screenshots that should be captured for submission."
    )
    add_table(
        doc,
        ["Evidence screenshot", "URL or command", "Expected result"],
        [
            ["Frontend SPA", "http://localhost:80", "BorderWatch login page and dashboard after login."],
            ["Gateway health", "http://localhost/health", "{\"status\":\"ok\"} response."],
            ["Grafana", "http://localhost:3000", "Grafana login with Prometheus datasource configured."],
            ["Prometheus targets", "http://localhost:9090/targets", "ingestion, tracking, compliance and node-exporter targets UP."],
            ["RabbitMQ management", "http://localhost:15672", "RabbitMQ management console for queues/exchange."],
            ["Container status", "docker compose ps", "All application containers running/healthy."],
            ["Smoke telemetry", "python simulator/simulator.py --rps 100 --duration 10", "Accepted telemetry requests and P95 latency output."],
        ],
        [1.8, 2.4, 3.0],
    )

    add_heading(doc, "11. Security Notes and Production Hardening", 1)
    add_callout(
        doc,
        "Important",
        "The repository currently supports a demonstration deployment. For production, rotate all .env secrets, restrict admin ports by source IP or VPN, add TLS on port 443, validate each GPS beacon API key against trucks.api_key_hash, and remove public demo credentials.",
    )
    add_bullets(
        doc,
        [
            "Keep only Nginx public; do not expose Spring service ports directly.",
            "Restrict Grafana, Prometheus, Alertmanager, RabbitMQ management, and MySQL to administrators.",
            "Replace default JWT, MySQL, RabbitMQ, and Grafana secrets before deployment.",
            "Use HTTPS with Let's Encrypt or another certificate authority for browser and beacon traffic.",
            "Add backup policy for MySQL volume and RabbitMQ state if deployed beyond a classroom demo.",
        ],
    )

    add_heading(doc, "12. Marking Summary", 1)
    add_table(
        doc,
        ["Requirement", "How BorderWatch satisfies it"],
        [
            ["Provision cloud/on-premise resources", "The stack is VPS-ready and deploys on a Docker-capable Linux host using docker-compose.yml."],
            ["Install necessary technologies", "Dockerfiles and Compose install Java services, Node/Vite frontend build, Nginx, MySQL, RabbitMQ and monitoring tools."],
            ["Configure networking/firewall/reverse proxies", "Nginx gateway routes API traffic; Docker bridge network isolates services; recommended UFW rules are documented."],
            ["Use Docker/Kubernetes/relevant tools", "Docker Compose is the implemented orchestration tool. Kubernetes/Terraform can be added later but are not required for the current repository."],
            ["Infrastructure diagram", "Included in this document as Figure 1."],
            ["Terraform/Bash scripts", "Bash scripts are present under scripts/ and CI automation is described in Jenkinsfile."],
            ["Screenshots and links", "Evidence links and screenshot targets are listed for frontend, health, Prometheus, Grafana, RabbitMQ and Compose status."],
        ],
        [2.2, 4.8],
    )

    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("BorderWatch - Section 1 Infrastructure Setup - Updated from project repository evidence")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
